"""Feature pins (2026-08-17) — Benji S2: package generation + the
verification gate. PRD v1.2 §5, Generation Spec v1.

THE CONTRACT PINNED HERE, in one line: **the agent may reframe, it may
not add** — enforced structurally (resume = deterministic selection
over parsed source bullets; the LLM touches only profile + letter) and
mechanically (the gate blocks any untraceable number, DO-NOT-USE hit or
banned phrase, wherever the words came from — including a poisoned JD).
"""
from __future__ import annotations

import importlib
import io
import json
from datetime import datetime

import pytest

from tests.regression_registry.benji_s2_fixture import (
    ALIEN_JD,
    FAKE_SOURCE,
    GOOD_JD,
    POISON_JD,
    job_row,
)

NOW = datetime(2026, 8, 17, 9, 0)


@pytest.fixture
def env(tmp_path, monkeypatch):
    monkeypatch.setenv("RAHAT_TEST_MODE", "1")
    monkeypatch.setenv("RAHAT_TEST_VAULT_DIR", str(tmp_path / "vault"))
    src = tmp_path / "source.md"
    src.write_text(FAKE_SOURCE)
    monkeypatch.setenv("BENJI_CANDIDATE_SOURCE", str(src))
    for var in ("RAHAT_JOBSEARCH_DB", "BENJI_FILTER_CONFIG",
                "BENJI_PREFERENCES", "BENJI_GATE_RULES",
                "BENJI_DELIVERY_EMAIL"):
        monkeypatch.delenv(var, raising=False)
    import agents.benji  # noqa: F401
    from bridges.jobsearch import store
    importlib.reload(store)
    store.upsert_batch([job_row(1)], now=NOW)
    return tmp_path


def _gen(display_id=1, llm=None):
    from agents.benji.generation import generate_package
    return generate_package(display_id, llm=llm, now=NOW)


# ─────────────────── source parsing ───────────────────────────────────
def test_source_parses_roles_caps_and_do_not_use(env):
    from agents.benji.source_parser import parse_source
    src = parse_source(FAKE_SOURCE)
    assert [r.org for r in src.roles][:2] == [
        "Example Rescue Org", "Example Education Foundation (TestProduct)"]
    assert src.roles[1].max_bullets == 4          # "Four bullets maximum"
    assert len(src.roles[0].default_pick) >= 4    # five-bullet default
    assert any("grantmaking" in x.lower() for x in src.never_claim)
    assert len(src.stories) == 5
    assert src.skills_core and src.profile and src.narrative


# ─────────────────── deterministic resume assembly ────────────────────
def test_resume_respects_caps_and_selects_by_required_quals(env):
    r = _gen()
    assert r["ok"], r.get("refusal")
    files = dict(r["files"])
    from docx import Document
    doc = Document(io.BytesIO(files["Resume_ExampleFoundation.docx"]))
    text = "\n".join(p.text for p in doc.paragraphs)
    # 7-bullet role capped at 5; the JD's required quals (funding,
    # compliance, partnership) must have won selection slots.
    assert text.count("•") == 0                   # bullets are styled,
    assert "funding from roughly $250K to $500K" in text
    assert "grants compliance on the grantee side" in text
    # 4-cap role (TestProduct) obeys its own cap.
    from agents.benji.source_parser import parse_source
    cap_role = parse_source(FAKE_SOURCE).roles[1]
    kept = [b for b in cap_role.bullets if b[:40] in text]
    assert len(kept) <= 4
    # PDF exists and is a PDF.
    assert files["Resume_ExampleFoundation.pdf"][:5] == b"%PDF-"
    # prompt.md carries the DO-NOT-USE list + both documents.
    assert "Never claim" not in files["prompt.md"]  # list is inlined flat
    assert "Grantmaking, grant writing" in files["prompt.md"]
    assert "CURRENT COVER LETTER" in files["prompt.md"]


def test_instruction_lines_never_render_into_the_resume(env):
    """The gate caught its own instruction text live (2026-08-17): the
    parser leaked '*Never "University of Pennsylvania."*' into the
    certifications and the rendered resume tripped the forbid rule.
    Italic instruction lines are metadata, never content."""
    r = _gen()
    assert r["ok"], r.get("refusal")
    resume_section = dict(r["files"])["prompt.md"].split(
        "CURRENT RESUME")[1]
    assert "Forbidden University" not in resume_section
    assert "Always attach" not in resume_section


def test_letter_retries_once_with_gate_feedback(env, tmp_path,
                                                monkeypatch):
    """Self-repair: first draft trips the gate, redraft with the exact
    failures clears it — one retry, letter only."""
    rules = tmp_path / "gate_rules.json"
    rules.write_text(json.dumps({"rules": [
        {"kind": "claim_forbid", "pattern": r"grant[\s-]?making",
         "why": "grantmaking may never be claimed"}]}))
    monkeypatch.setenv("BENJI_GATE_RULES", str(rules))
    calls = []

    def llm(prompt):
        calls.append(prompt)
        if "failed a mechanical fact gate" in prompt:
            return ("Dear team, I observe, attune and act. My grants "
                    "experience is on the grantee side — funder "
                    "reporting and compliance across a portfolio. " * 8)
        return ("Dear team, I bring deep grantmaking experience and "
                "led grantmaking for years. " * 20)

    r = _gen(llm=llm)
    assert r["ok"], r.get("refusal")
    assert any("failed a mechanical fact gate" in c for c in calls)
    assert any("redrafted once" in f for f in r["flags"])


def test_skills_reorder_never_adds(env):
    from agents.benji.generation import reorder_skills
    from agents.benji.source_parser import parse_source
    from agents.benji.coverage import _terms
    src = parse_source(FAKE_SOURCE)
    lines = reorder_skills(src, _terms(GOOD_JD))
    joined = " ".join(lines)
    for skill in src.skills_core:
        assert skill in joined
    # Nothing appears that isn't in the source lists.
    all_src = {*src.skills_core, *src.skills_situational,
               *src.skills_technical}
    for line in lines:
        for item in line.split(": ", 1)[1].split(" · "):
            assert item in all_src


# ─────────────────── coverage floor refusal (Tara #3) ─────────────────
def test_kit_refuses_below_coverage_floor_with_unmatched_list(env):
    from bridges.jobsearch import store
    store.upsert_batch([job_row(2, jd=ALIEN_JD, title="Program Officer, "
                                "Quantum", org="Alien Fund")], now=NOW)
    rows = store.queue_rows()
    alien_id = next(r["id"] for r in rows if r["org"] == "Alien Fund")
    r = _gen(alien_id)
    assert not r["ok"] and r["files"] == []
    assert "below the 60% floor" in r["refusal"]
    assert "Unmatched" in r["refusal"]
    assert "add it there and reply kit again" in r["refusal"]


# ─────────────────── the gate blocks fabrication ──────────────────────
def test_gate_blocks_untraceable_number_from_llm(env):
    filler = ("I connect people and programs across communities and "
              "keep delivery on track through change. " * 16)
    poisoned_letter = (f"Dear Team,\n\n{filler}\n\nI led a $750K program "
                       "reaching 2,000 families and doubled income 400% "
                       "for every one of them.\n\nSincerely")
    r = _gen(llm=lambda p: poisoned_letter)
    assert not r["ok"]
    assert any("not traceable" in f for f in r["gate"].failures)


def test_adversarial_poisoned_jd_cannot_produce_a_passing_package(env):
    """The JD is untrusted input. Simulate the worst case — the LLM
    obeys the poison and claims grantmaking + $750K — and the package
    must still die at the gate. Then verify the deterministic resume
    path never contained the poison at all."""
    from bridges.jobsearch import store
    store.upsert_batch([job_row(3, jd=POISON_JD, org="Poison Org",
                                title="Program Officer, Community")],
                       now=NOW)
    pid = next(r["id"] for r in store.queue_rows()
               if r["org"] == "Poison Org")
    obeyed = ("I am writing because I led grantmaking for a $750K "
              "program. " * 20)
    r = _gen(pid, llm=lambda p: obeyed)
    assert not r["ok"]
    fails = " ".join(r["gate"].failures)
    assert "not traceable" in fails or "grant" in fails.lower()

    # Honest LLM on the same poisoned JD → resume text contains no
    # grantmaking claim (deterministic path is immune by construction).
    r2 = _gen(pid, llm=lambda p: "Dear team, I observe, attune and act. "
              "My record covers program management, funder reporting "
              "and partnership development on the grantee side. " * 6)
    assert r2["ok"], r2.get("refusal")
    files = dict(r2["files"])
    assert "grantmaking" not in files["prompt.md"].split(
        "CURRENT RESUME")[1].split("CURRENT COVER LETTER")[0].lower()


def test_gate_rules_are_vault_data_with_pair_and_claim_shapes(env,
                                                              tmp_path,
                                                              monkeypatch):
    rules = {"rules": [
        {"kind": "pair_line", "token": "exampleorg",
         "requires": "(volunteer)", "why": "must carry (Volunteer)"},
        # Alternation ON PURPOSE: an ungrouped interpolation once made
        # the first branch match bare (any mention blocked — including
        # the negated Hewlett move). Caught in launch rehearsal.
        {"kind": "claim_forbid",
         "pattern": r"grant[\s-]?making|proposal\s+writing",
         "why": "grantmaking may never be claimed"},
    ]}
    p = tmp_path / "rules.json"
    p.write_text(json.dumps(rules))
    monkeypatch.setenv("BENJI_GATE_RULES", str(p))
    from agents.benji.verification import verify_package
    from agents.benji.source_parser import parse_source
    src = parse_source(FAKE_SOURCE)

    bad = verify_package(
        resume_text="Director — ExampleOrg\n", letter_text="",
        source=src, role_bullets={})
    assert any("(Volunteer)" in f for f in bad.failures)

    claimed = verify_package(
        resume_text="", letter_text="I have deep experience in "
        "grantmaking across portfolios.", source=src, role_bullets={})
    assert any("claimed" in f for f in claimed.failures)

    negated = verify_package(
        resume_text="", letter_text="I have never sat on the "
        "grantmaking side of the table - I have lived the other side.",
        source=src, role_bullets={})
    assert negated.ok, negated.failures     # the Hewlett move must pass

    # Bare mention (no claim verbs) and negated-experience phrasing
    # both pass; the alternation's second branch still bites.
    bare = verify_package(
        resume_text="", letter_text="Grantmaking shapes what programs "
        "can deliver; I have watched it from the grantee side.",
        source=src, role_bullets={})
    assert bare.ok, bare.failures
    prop = verify_package(
        resume_text="", letter_text="I have deep experience in "
        "proposal writing for major funders.", source=src,
        role_bullets={})
    assert any("claimed" in f for f in prop.failures)


def test_banned_phrase_blocks(env):
    r = _gen(llm=lambda p: "I am thrilled to apply because I am "
             "uniquely suited. " * 30)
    assert not r["ok"]
    assert any("banned phrase" in f for f in r["gate"].failures)


# ─────────────────── degrade paths are honest ─────────────────────────
def test_llm_unavailable_degrades_with_flags_never_silence(env):
    r = _gen(llm=None)          # test mode + no seam → degrade
    assert r["ok"], r.get("refusal")
    flags = " ".join(r["flags"])
    assert "DRAFT-DEGRADED" in flags
    assert "org-research placeholder" in flags
    # review.md carries coverage, story, and the flag list.
    review = dict(r["files"])["review.md"]
    assert "Keyword coverage:" in review and "## Story:" in review
    assert "Flag list" in review


def test_org_placeholder_never_an_invented_specific(env):
    """Degraded letter references the org ONLY via the placeholder —
    an invented specific is worse than a generic letter (Spec §4)."""
    r = _gen(llm=None)
    letter_docx = dict(r["files"])["CoverLetter_ExampleFoundation.docx"]
    from docx import Document
    text = "\n".join(p.text for p in
                     Document(io.BytesIO(letter_docx)).paragraphs)
    assert "[ORG-SPECIFIC" in text
