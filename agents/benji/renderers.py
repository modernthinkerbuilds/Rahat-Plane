"""benji.renderers — one resume/letter model, two renderers (docx + PDF).

ATS-safe by construction (Generation Spec §2): single column, no
tables, no text boxes, no headers/footers carrying content, standard
fonts, and no role entry split across a PDF page (KeepTogether). The
docx is the submission primary; the PDF renders alongside for her
review. Both are byte-deterministic given the same model + timestamps.

COORDINATION note: introduces python-docx + reportlab (requirements.txt
change, announced).
"""
from __future__ import annotations

import io
from dataclasses import dataclass, field


@dataclass
class ResumeModel:
    name: str
    contact: str
    location: str
    profile: str
    roles: list[dict] = field(default_factory=list)
    # role dict: {title, org, dates, location, bullets[], early: bool}
    education: list[str] = field(default_factory=list)
    certifications: list[str] = field(default_factory=list)
    skills_lines: list[str] = field(default_factory=list)
    volunteer: list[str] = field(default_factory=list)

    def as_text(self) -> str:
        """The gate greps this — it must contain every word the reader
        will see."""
        parts = [self.name, self.contact, self.location, "", self.profile]
        for r in self.roles:
            parts += ["", f"{r['title']} — {r['org']}",
                      " · ".join(x for x in (r.get("dates", ""),
                                             r.get("location", "")) if x)]
            parts += [f"• {b}" for b in r["bullets"]]
        parts += ["", *self.education, *self.certifications,
                  *self.skills_lines, *self.volunteer]
        return "\n".join(p for p in parts if p is not None)


def render_resume_docx(m: ResumeModel) -> bytes:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    style = d.styles["Normal"]
    style.font.name, style.font.size = "Calibri", Pt(10.5)

    h = d.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(m.name)
    run.bold, run.font.size = True, Pt(16)
    c = d.add_paragraph(f"{m.contact} · {m.location}")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER

    d.add_paragraph(m.profile)

    def heading(text: str):
        p = d.add_paragraph()
        r = p.add_run(text.upper())
        r.bold, r.font.size = True, Pt(11.5)

    heading("Experience")
    for r in m.roles:
        p = d.add_paragraph()
        rr = p.add_run(f"{r['title']} — {r['org']}")
        rr.bold = True
        meta = " · ".join(x for x in (r.get("dates", ""),
                                      r.get("location", "")) if x)
        if meta:
            mr = p.add_run(f"   {meta}")
            mr.italic = True
        for b in r["bullets"]:
            d.add_paragraph(b, style="List Bullet")

    if m.education or m.certifications:
        heading("Education & Certifications")
        for line in (*m.education, *m.certifications):
            d.add_paragraph(line)
    if m.skills_lines:
        heading("Skills")
        for line in m.skills_lines:
            d.add_paragraph(line)
    if m.volunteer:
        heading("Volunteer")
        for line in m.volunteer:
            d.add_paragraph(line)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _pdf_story(m: ResumeModel):
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import KeepTogether, Paragraph, Spacer

    ss = getSampleStyleSheet()
    base = ParagraphStyle("base", parent=ss["Normal"], fontSize=9.5,
                          leading=12)
    name_s = ParagraphStyle("n", parent=base, fontSize=15, leading=18,
                            alignment=1, spaceAfter=2)
    head_s = ParagraphStyle("h", parent=base, fontSize=10.5,
                            spaceBefore=8, spaceAfter=2)
    story = [Paragraph(f"<b>{m.name}</b>", name_s),
             Paragraph(f"{m.contact} · {m.location}",
                       ParagraphStyle("c", parent=base, alignment=1)),
             Spacer(1, 6), Paragraph(m.profile, base), Spacer(1, 4),
             Paragraph("<b>EXPERIENCE</b>", head_s)]
    for r in m.roles:
        meta = " · ".join(x for x in (r.get("dates", ""),
                                      r.get("location", "")) if x)
        block = [Paragraph(f"<b>{r['title']} — {r['org']}</b>"
                           + (f"  <i>{meta}</i>" if meta else ""), base)]
        block += [Paragraph(f"•  {b}", base) for b in r["bullets"]]
        block.append(Spacer(1, 4))
        story.append(KeepTogether(block))   # never split a role (§2)
    for title, lines in (("EDUCATION & CERTIFICATIONS",
                          [*m.education, *m.certifications]),
                         ("SKILLS", m.skills_lines),
                         ("VOLUNTEER", m.volunteer)):
        if lines:
            story.append(Paragraph(f"<b>{title}</b>", head_s))
            story += [Paragraph(x, base) for x in lines]
    return story


def render_resume_pdf(m: ResumeModel) -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER, leftMargin=0.7 * inch,
                            rightMargin=0.7 * inch, topMargin=0.55 * inch,
                            bottomMargin=0.55 * inch, title="Resume")
    doc.build(_pdf_story(m))
    return buf.getvalue()


def render_letter_docx(name: str, contact: str, body: str) -> bytes:
    from docx import Document
    from docx.shared import Pt

    d = Document()
    d.styles["Normal"].font.name = "Calibri"
    d.styles["Normal"].font.size = Pt(11)
    p = d.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    d.add_paragraph(contact)
    d.add_paragraph("")
    for para in body.split("\n\n"):
        if para.strip():
            d.add_paragraph(para.strip())
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def render_letter_pdf(name: str, contact: str, body: str) -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    ss = getSampleStyleSheet()
    base = ParagraphStyle("b", parent=ss["Normal"], fontSize=10.5,
                          leading=14)
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER, leftMargin=0.9 * inch,
                            rightMargin=0.9 * inch, topMargin=0.8 * inch,
                            bottomMargin=0.8 * inch, title="Cover Letter")
    story = [Paragraph(f"<b>{name}</b>", base), Paragraph(contact, base),
             Spacer(1, 12)]
    for para in body.split("\n\n"):
        if para.strip():
            story += [Paragraph(para.strip(), base), Spacer(1, 8)]
    doc.build(story)
    return buf.getvalue()
